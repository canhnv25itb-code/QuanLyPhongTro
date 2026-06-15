import os
import sys
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT = "BaoCao_QuanLyPhongTro.docx"

# ==========================================
# DIAGRAM GENERATION SECTION (Pillow)
# ==========================================

def get_pil_font(font_name="arial.ttf", size=14, bold=False):
    win_font_dir = "C:\\Windows\\Fonts"
    filename = "arialbd.ttf" if bold else "arial.ttf"
    path = os.path.join(win_font_dir, filename)
    if os.path.exists(path):
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()

def draw_usecase_diagram(output_path):
    img = Image.new("RGB", (1000, 700), "white")
    draw = ImageDraw.Draw(img)
    
    font_title = get_pil_font("arial", 20, bold=True)
    font_bold = get_pil_font("arial", 13, bold=True)
    font_normal = get_pil_font("arial", 11, bold=False)
    
    # Title
    draw.text((500, 30), "SƠ ĐỒ USE CASE HỆ THỐNG QUẢN LÝ PHÒNG TRỌ", fill="black", font=font_title, anchor="mm")
    
    # System boundary
    draw.rectangle([220, 80, 780, 650], outline="#2C3E50", width=3)
    draw.text((500, 95), "Hệ Thống Quản Lý Phòng Trọ (Desktop App)", fill="#2C3E50", font=font_bold, anchor="mm")
    
    # Draw Actor 1: Chủ Trọ (Left)
    # Head
    draw.ellipse([80, 280, 120, 320], outline="#1A5276", width=3, fill="#AED6F1")
    # Body
    draw.line([100, 320, 100, 400], fill="#1A5276", width=3)
    # Arms
    draw.line([70, 345, 130, 345], fill="#1A5276", width=3)
    # Legs
    draw.line([100, 400, 75, 460], fill="#1A5276", width=3)
    draw.line([100, 400, 125, 460], fill="#1A5276", width=3)
    draw.text((100, 475), "Chủ trọ\n(Landlord)", fill="#1A5276", font=font_bold, anchor="mm", align="center")
    
    # Draw Actor 2: Khách Thuê (Right)
    # Head
    draw.ellipse([880, 280, 920, 320], outline="#27AE60", width=3, fill="#ABEBC6")
    # Body
    draw.line([900, 320, 900, 400], fill="#27AE60", width=3)
    # Arms
    draw.line([870, 345, 930, 345], fill="#27AE60", width=3)
    # Legs
    draw.line([900, 400, 875, 460], fill="#27AE60", width=3)
    draw.line([900, 400, 925, 460], fill="#27AE60", width=3)
    draw.text((900, 475), "Khách thuê\n(Tenant)", fill="#27AE60", font=font_bold, anchor="mm", align="center")
    
    # Use cases
    use_cases = [
        (360, 140, "Đăng nhập &\nĐăng ký tài khoản", True, True),
        (360, 205, "Quản lý phòng trọ\n(Thêm, sửa, xóa, tìm)", True, False),
        (360, 270, "Quản lý khách thuê\n(Liên kết tài khoản)", True, False),
        (360, 335, "Lập & quản lý hợp đồng\n(Tạo hợp đồng mặc định)", True, False),
        (360, 400, "Nhập số điện nước\n(Đồng bộ theo tháng)", True, False),
        (360, 465, "Quản lý phí dịch vụ\n(Wifi, xe, vệ sinh...)", True, False),
        (360, 530, "Lập & quản lý hóa đơn\n(Tự động tính tiền)", True, False),
        (360, 595, "Xem Dashboard &\nBiểu đồ thống kê", True, False),
        (640, 360, "Xem thông tin phòng,\nDịch vụ & Điện nước", False, True),
        (640, 430, "Xem & Theo dõi hóa đơn\n(Trạng thái thanh toán)", False, True),
    ]
    
    for x, y, text, to_chu_tro, to_khach_thue in use_cases:
        rx, ry = 110, 25
        draw.ellipse([x - rx, y - ry, x + rx, y + ry], outline="#2980B9", fill="#EBF5FB", width=2)
        draw.text((x, y), text, fill="#2C3E50", font=font_normal, anchor="mm", align="center")
        
        if to_chu_tro:
            draw.line([130, 350, x - rx, y], fill="#7F8C8D", width=1)
        if to_khach_thue:
            draw.line([870, 350, x + rx, y], fill="#7F8C8D", width=1)

    img.save(output_path, "PNG")
    print(f"Use Case diagram generated at {output_path}")

def draw_architecture_diagram(output_path):
    img = Image.new("RGB", (1000, 700), "white")
    draw = ImageDraw.Draw(img)
    
    font_title = get_pil_font("arial", 20, bold=True)
    font_layer = get_pil_font("arial", 14, bold=True)
    font_bold = get_pil_font("arial", 12, bold=True)
    font_normal = get_pil_font("arial", 11, bold=False)
    
    # Title
    draw.text((500, 30), "SƠ ĐỒ KIẾN TRÚC PHẦN MỀM (3-TIER ARCHITECTURE)", fill="black", font=font_title, anchor="mm")
    
    def draw_box(x1, y1, x2, y2, title, items, fill_color, outline_color):
        draw.rectangle([x1, y1, x2, y2], fill=fill_color, outline=outline_color, width=2)
        draw.rectangle([x1, y1, x2, y1 + 30], fill=outline_color)
        draw.text(((x1+x2)//2, y1 + 15), title, fill="white", font=font_bold, anchor="mm")
        
        y_offset = y1 + 45
        for item in items:
            draw.text((x1 + 15, y_offset), "• " + item, fill="#2C3E50", font=font_normal)
            y_offset += 20
            
    # Tier 1: UI Layer
    draw_box(100, 80, 900, 200, 
             "LỚP GIAO DIỆN (UI LAYER - PRESENTATION)", 
             ["MainFrame & KhachThueFrame: Khung giao diện chính điều hướng ứng dụng",
              "LoginForm & RegisterForm: Xử lý đăng nhập, đăng ký tài khoản",
              "Các Panel nghiệp vụ: PhongTroPanel, KhachThuePanel, HopDongPanel, DienNuocPanel, DichVuPanel, HoaDonPanel",
              "Các Dialog nhập liệu: PhongTroDialog, KhachThueDialog, v.v."], 
             "#FEF9E7", "#F39C12")
             
    # Tier 2: DAO Layer
    draw_box(100, 260, 900, 390, 
             "LỚP TRUY XUẤT DỮ LIỆU (DAO LAYER - DATA ACCESS)", 
             ["TaiKhoanDAO & PhongTroDAO: Truy vấn xác thực và quản lý phòng trọ",
              "KhachThueDAO & HopDongDAO: Xử lý thông tin khách thuê và hợp đồng",
              "DienNuocDAO & DichVuDAO: Ghi nhận chỉ số điện nước, phí dịch vụ",
              "HoaDonDAO & ThongKeDAO: Tính hóa đơn, truy xuất doanh thu báo cáo cho dashboard"], 
             "#E8F8F5", "#16A085")
             
    # Tier 3: Model & DB Layer
    draw_box(100, 450, 480, 600, 
             "LỚP ĐỐI TƯỢNG (MODEL CLASSES)", 
             ["TaiKhoan, PhongTro, KhachThue",
              "HopDong, HoaDon, DienNuoc, DichVu",
              "Lưu trữ dữ liệu dạng đối tượng Java",
              "Để truyền tải dữ liệu giữa UI và DAO"], 
             "#EBF5FB", "#2980B9")
             
    draw_box(520, 450, 900, 600, 
             "CƠ SỞ DỮ LIỆU (MS SQL SERVER)", 
             ["Database: QuanLyPhongTro",
              "Chứa các bảng dữ liệu có ràng buộc toàn vẹn",
              "Truy xuất thông qua Microsoft JDBC Driver",
              "Quản lý theo phiên làm việc (Session)"], 
             "#F2F3F4", "#7F8C8D")
             
    # Tier 4: Util Layer
    draw.rectangle([100, 620, 900, 670], fill="#FDEDEC", outline="#C0392B", width=2)
    draw.text((500, 645), "CÁC TIỆN ÍCH HỆ THỐNG (UTIL): DatabaseConnection | DatabaseMigration (Đồng bộ schema) | Session", 
              fill="#C0392B", font=font_bold, anchor="mm")
              
    # Arrows
    draw.line([500, 200, 500, 260], fill="#E67E22", width=3)
    draw.polygon([495, 255, 505, 255, 500, 260], fill="#E67E22")
    draw.text((515, 230), "Gọi DAO xử lý", fill="#E67E22", font=font_bold, anchor="lm")
    
    draw.line([710, 390, 710, 450], fill="#16A085", width=3)
    draw.polygon([705, 445, 715, 445, 710, 450], fill="#16A085")
    draw.text((725, 420), "SQL Query / JDBC", fill="#16A085", font=font_bold, anchor="lm")
    
    draw.line([290, 390, 290, 450], fill="#2980B9", width=3)
    draw.polygon([285, 445, 295, 445, 290, 450], fill="#2980B9")
    draw.text((200, 420), "Trả về Model Objects", fill="#2980B9", font=font_bold, anchor="rm")

    img.save(output_path, "PNG")
    print(f"Architecture diagram generated at {output_path}")

def draw_erd_diagram(output_path):
    img = Image.new("RGB", (1100, 850), "white")
    draw = ImageDraw.Draw(img)
    
    font_title = get_pil_font("arial", 20, bold=True)
    font_bold = get_pil_font("arial", 12, bold=True)
    font_normal = get_pil_font("arial", 10, bold=False)
    
    # Title
    draw.text((550, 30), "SƠ ĐỒ CƠ SỞ DỮ LIỆU CHI TIẾT (ENTITY RELATIONSHIP DIAGRAM)", fill="black", font=font_title, anchor="mm")
    
    def draw_table(x, y, w, title, pks, fks, normal_cols):
        h = 35 + 20 * (len(pks) + len(fks) + len(normal_cols))
        draw.rectangle([x, y, x + w, y + h], fill="#F8F9F9", outline="#34495E", width=2)
        draw.rectangle([x, y, x + w, y + 30], fill="#34495E")
        draw.text((x + w//2, y + 15), title, fill="white", font=font_bold, anchor="mm")
        
        curr_y = y + 45
        for pk in pks:
            draw.text((x + 10, curr_y), "🔑 " + pk, fill="#C0392B", font=font_bold)
            curr_y += 20
        for fk in fks:
            draw.text((x + 10, curr_y), "🔗 " + fk, fill="#2980B9", font=font_bold)
            curr_y += 20
        for col in normal_cols:
            draw.text((x + 10, curr_y), "  " + col, fill="#2C3E50", font=font_normal)
            curr_y += 20
        return x, y, w, h
        
    # Row 1
    t_tk = draw_table(50, 100, 260, "TaiKhoan", 
                      ["TenDN: VARCHAR(50)"], 
                      [], 
                      ["MatKhau: VARCHAR(50)", "Quyen: NVARCHAR(50)", "Email: VARCHAR(100)"])
                      
    t_pt = draw_table(420, 100, 260, "PhongTro", 
                      ["MaPhong: VARCHAR(50)", "TenDN: VARCHAR(50)"], 
                      [], 
                      ["Tang: INT", "TrangThai: NVARCHAR(50)", "GiaPhong: FLOAT", "DienTich: FLOAT"])
                      
    t_dv = draw_table(790, 100, 260, "DichVu", 
                      ["MaDV: INT (Identity)"], 
                      ["MaPhong: VARCHAR(50)", "TenDN: VARCHAR(50)"], 
                      ["Wifi: FLOAT", "ThangMay: FLOAT", "GiuXe: FLOAT", "PCCC_VeSinh: FLOAT", "TongTienDV: FLOAT"])

    # Row 2
    t_hd = draw_table(50, 380, 260, "HopDong", 
                      ["MaHD: INT (Identity)"], 
                      ["MaPhong: VARCHAR(50)", "TenDN: VARCHAR(50)"], 
                      ["TenKhachHang: NVARCHAR(100)", "NgayBatDau: DATE", "NgayKetThuc: DATE", "TienCoc: FLOAT", "TrangThai: NVARCHAR(50)"])
                      
    t_kt = draw_table(420, 380, 260, "KhachThue", 
                      ["MaKhach: INT (Identity)"], 
                      ["MaPhong: VARCHAR(50)", "TenDN: VARCHAR(50)", "MaChuTro: VARCHAR(50)"], 
                      ["HoTen: NVARCHAR(100)", "CCCD: VARCHAR(20)", "SDT: VARCHAR(15)", "QueQuan: NVARCHAR(200)", "TrangThai: NVARCHAR(50)"])
                      
    t_dn = draw_table(790, 380, 260, "DienNuoc", 
                      ["ID: INT (Identity)"], 
                      ["MaPhong: VARCHAR(50)", "TenDN: VARCHAR(50)"], 
                      ["ThangNam: VARCHAR(10)", "DienDau: INT", "DienCuoi: INT", "NuocDau: INT", "NuocCuoi: INT", "TienDien: FLOAT", "TienNuoc: FLOAT"])

    # Row 3
    t_hdon = draw_table(420, 650, 260, "HoaDon", 
                        ["MaHoaDon: INT (Identity)"], 
                        ["TenPhong: VARCHAR(50)", "TenDN: VARCHAR(50)"], 
                        ["NgayLap: DATE", "TienPhong: FLOAT", "TienDien: FLOAT", "TienNuoc: FLOAT", "TienDichVu: FLOAT", "TrangThai: NVARCHAR(50)"])

    # Relations
    draw.line([310, 150, 420, 150], fill="#7F8C8D", width=2)
    draw.text((320, 130), "1", fill="black", font=font_bold)
    draw.text((400, 130), "N", fill="black", font=font_bold)
    
    draw.line([680, 150, 790, 150], fill="#7F8C8D", width=2)
    draw.text((690, 130), "1", fill="black", font=font_bold)
    draw.text((770, 130), "1", fill="black", font=font_bold)

    draw.line([550, 280, 550, 380], fill="#7F8C8D", width=2)
    draw.text((560, 290), "1", fill="black", font=font_bold)
    draw.text((560, 360), "N", fill="black", font=font_bold)

    draw.line([420, 180, 350, 180], fill="#7F8C8D", width=2)
    draw.line([350, 180, 350, 420], fill="#7F8C8D", width=2)
    draw.line([350, 420, 310, 420], fill="#7F8C8D", width=2)
    draw.text((400, 160), "1", fill="black", font=font_bold)
    draw.text((320, 425), "N", fill="black", font=font_bold)

    draw.line([680, 220, 740, 220], fill="#7F8C8D", width=2)
    draw.line([740, 220, 740, 420], fill="#7F8C8D", width=2)
    draw.line([740, 420, 790, 420], fill="#7F8C8D", width=2)
    draw.text((690, 200), "1", fill="black", font=font_bold)
    draw.text((770, 425), "N", fill="black", font=font_bold)

    draw.line([420, 240, 390, 240], fill="#7F8C8D", width=2)
    draw.line([390, 240, 390, 700], fill="#7F8C8D", width=2)
    draw.line([390, 700, 420, 700], fill="#7F8C8D", width=2)
    draw.text((405, 220), "1", fill="black", font=font_bold)
    draw.text((405, 705), "N", fill="black", font=font_bold)

    draw.line([180, 280, 180, 340], fill="#7F8C8D", width=2)
    draw.line([180, 340, 450, 340], fill="#7F8C8D", width=2)
    draw.line([450, 340, 450, 380], fill="#7F8C8D", width=2)
    draw.text((190, 290), "1", fill="black", font=font_bold)
    draw.text((460, 360), "N", fill="black", font=font_bold)

    img.save(output_path, "PNG")
    print(f"ERD diagram generated at {output_path}")


# ==========================================
# REPORT FORMATTING HELPERS (python-docx)
# ==========================================

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text(cell, text, bold=False, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def style_table(table, widths, header_fill="F2F4F7"):
    set_table_widths(table, widths)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.font.color.rgb = None
    set_repeat_table_header(table.rows[0])
    
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", 80), ("bottom", 80), ("left", 120), ("right", 120)]:
                elem = tc_mar.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    tc_mar.append(elem)
                elem.set(qn("w:w"), str(value))
                elem.set(qn("w:type"), "dxa")

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20") # 10pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

def add_toc(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    fld.set(qn("w:dirty"), "true")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Mục lục sẽ được Word tự động cập nhật khi mở tài liệu."
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

def enable_update_fields(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

def add_run_tnr(p, text, size=13, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r

def add_paragraph_tnr(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    if level == 1:
        p.paragraph_format.keep_with_next = True
        text = text.upper()
        
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(16 if level == 1 else (14 if level == 2 else 13))
    r.bold = True
    if level == 3:
        r.font.italic = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        
        r_bullet = p.add_run("–  ")
        r_bullet.font.name = "Times New Roman"
        r_bullet.font.size = Pt(13)
        r_bullet.bold = True
        
        r_text = p.add_run(item)
        r_text.font.name = "Times New Roman"
        r_text.font.size = Pt(13)

def add_numbered(doc, items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        
        r_num = p.add_run(f"{idx}. ")
        r_num.font.name = "Times New Roman"
        r_num.font.size = Pt(13)
        r_num.bold = True
        
        r_text = p.add_run(item)
        r_text.font.name = "Times New Roman"
        r_text.font.size = Pt(13)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_picture(doc, path, width_inches):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=Inches(width_inches))
    return p

def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Nội dung", True)
    set_cell_text(table.rows[0].cells[1], "Mô tả", True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    style_table(table, [1.9, 4.5])
    return table

def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after in [
        ("Heading 1", 16, 14, 6),
        ("Heading 2", 14, 10, 6),
        ("Heading 3", 13, 6, 4),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    styles["Caption"].font.name = "Times New Roman"
    styles["Caption"].font.size = Pt(11)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor(0, 0, 0)


def add_cover_page_border(section):
    sectPr = section._sectPr
    pgBorders = sectPr.find(qn("w:pgBorders"))
    if pgBorders is not None:
        sectPr.remove(pgBorders)
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    pgBorders.set(qn("w:display"), "firstPage")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "12") # 1.5 pt
        border.set(qn("w:space"), "24") # margin in points from page edge
        border.set(qn("w:color"), "auto")
        pgBorders.append(border)
    sectPr.append(pgBorders)


def cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0)
    
    add_run_tnr(p, "ĐẠI HỌC ĐÀ NẴNG\n", 13, bold=True)
    add_run_tnr(p, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ VÀ TRUYỀN THÔNG VIỆT - HÀN\n", 13, bold=True)
    add_run_tnr(p, "KHOA CÔNG NGHỆ THÔNG TIN\n", 13, bold=True)
    add_run_tnr(p, "~~~~~~~*~~~~~~~\n\n", 12, bold=True)

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(24)
    p_logo.paragraph_format.first_line_indent = Inches(0)
    run_logo = p_logo.add_run()
    run_logo.add_picture("vku_logo.png", width=Inches(1.5))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.first_line_indent = Inches(0)
    add_run_tnr(p_title, "BÁO CÁO ĐỒ ÁN MÔN HỌC\n", 22, bold=True)
    
    p_title_sub = doc.add_paragraph()
    p_title_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title_sub.paragraph_format.space_before = Pt(6)
    p_title_sub.paragraph_format.space_after = Pt(12)
    p_title_sub.paragraph_format.first_line_indent = Inches(0)
    add_run_tnr(p_title_sub, "ĐỀ TÀI: PHẦN MỀM QUẢN LÝ PHÒNG TRỌ\n", 16, bold=True)
    add_run_tnr(p_title_sub, "Hệ thống Quản lý cơ sở lưu trú và dịch vụ phòng trọ\n", 13, italic=True)

    p_tech = doc.add_paragraph()
    p_tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tech.paragraph_format.space_before = Pt(6)
    p_tech.paragraph_format.space_after = Pt(36)
    p_tech.paragraph_format.first_line_indent = Inches(0)
    add_run_tnr(p_tech, "Nền tảng phát triển: JAVA SWING, SQL SERVER, JDBC\n", 12, bold=True)

    # Info table
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info = [
        ("GIẢNG VIÊN HƯỚNG DẪN", ": Th.S Nguyễn Đăng Ý"),
        ("SINH VIÊN THỰC HIỆN",  ": Nguyễn Văn Cảnh"),
        ("MÃ SINH VIÊN",         ": 25ITB012"),
        ("SINH VIÊN THỰC HIỆN",  ": Nguyễn Nhật Nguyên"),
        ("MÃ SINH VIÊN",         ": 25ITB110"),
        ("LỚP",                  ": 25ITB"),
    ]
    
    for key, val in info:
        row = table.add_row()
        
        cell_space = row.cells[0]
        cell_space.text = ""
        p_space = cell_space.paragraphs[0]
        p_space.paragraph_format.first_line_indent = Inches(0)
        p_space.paragraph_format.space_after = Pt(0)
        
        cell_key = row.cells[1]
        p_key = cell_key.paragraphs[0]
        p_key.paragraph_format.first_line_indent = Inches(0)
        p_key.paragraph_format.space_after = Pt(0)
        add_run_tnr(p_key, key, 12, bold=True)
        
        cell_val = row.cells[2]
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.first_line_indent = Inches(0)
        p_val.paragraph_format.space_after = Pt(0)
        add_run_tnr(p_val, val, 12, bold=False)
        
    set_table_widths(table, [1.5, 2.5, 2.8])
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", 60), ("bottom", 60), ("left", 60), ("right", 60)]:
                elem = tc_mar.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    tc_mar.append(elem)
                elem.set(qn("w:w"), str(value))
                elem.set(qn("w:type"), "dxa")

    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.space_before = Pt(60)
    p_year.paragraph_format.space_after = Pt(0)
    p_year.paragraph_format.first_line_indent = Inches(0)
    add_run_tnr(p_year, "ĐÀ NẴNG – 2026", 12, bold=True)
    
    add_cover_page_border(doc.sections[0])
    
    # New section for the rest of document
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.footer.is_linked_to_previous = False
    
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(2.0)
    new_section.bottom_margin = Cm(2.0)
    new_section.left_margin = Cm(3.0)
    new_section.right_margin = Cm(2.0)
    new_section.header_distance = Cm(1.25)
    new_section.footer_distance = Cm(1.25)
    
    add_page_number(new_section.footer.paragraphs[0])
    
    sectPr1 = new_section._sectPr
    pgBorders1 = OxmlElement("w:pgBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        pgBorders1.append(border)
    sectPr1.append(pgBorders1)


def front_matter(doc):
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_paragraph_tnr(doc,
        "Em xin gửi lời cảm ơn chân thành đến quý thầy cô khoa Công nghệ thông tin trường Đại học Công nghệ thông tin và Truyền thông Việt - Hàn "
        "đã hướng dẫn, góp ý và tạo điều kiện thuận lợi nhất để em hoàn thành báo cáo đồ án môn học "
        "“Phần mềm quản lý phòng trọ”. Trong suốt quá trình thực hiện đồ án này, em đã có cơ hội củng cố lại các kiến thức quan trọng về lập trình Java, "
        "thiết kế các thành phần giao diện Java Swing, thao tác kết nối cơ sở dữ liệu Microsoft SQL Server thông qua thư viện kết nối JDBC và tổ chức "
        "cấu trúc mã nguồn dự án theo mô hình nhiều lớp phân rã."
    )
    add_paragraph_tnr(doc,
        "Em cũng xin gửi lời biết ơn tới gia đình, bạn bè và những người đã hỗ trợ, đóng góp ý kiến nghiệp vụ thực tế về hoạt động quản lý vận hành "
        "nhà trọ, giúp em kiểm thử tốt các chức năng phần mềm và hoàn thiện nội dung báo cáo đồ án môn học. Do kinh nghiệm phát triển phần mềm thực tế "
        "của bản thân còn hạn chế, nội dung báo cáo khó tránh khỏi những điểm thiếu sót; em rất mong nhận được những nhận xét, ý kiến đóng góp quý báu "
        "của quý thầy cô để sản phẩm đồ án được hoàn thiện tốt hơn."
    )

    add_heading(doc, "LỜI MỞ ĐẦU", 1)
    add_paragraph_tnr(doc,
        "Hoạt động vận hành quản lý các cơ sở lưu trú và phòng trọ hiện nay thường bao gồm rất nhiều nghiệp vụ lặp đi lặp lại hàng ngày như: quản lý tình trạng "
        "phòng trống, tiếp nhận hồ sơ đăng ký của khách thuê, lập hợp đồng thuê phòng pháp lý, ghi nhận chỉ số điện nước sử dụng hàng tháng, áp dụng các loại phí "
        "dịch vụ cố định, kết toán hóa đơn thu tiền và theo dõi số liệu doanh thu. Nếu các nghiệp vụ trên được xử lý theo cách thủ công truyền thống thông qua "
        "sổ ghi chép hoặc bảng tính Excel, người quản lý rất dễ gặp phải các sai sót số liệu tính toán điện nước, mất thời gian khi cần đối soát công nợ hoặc tra cứu "
        "lại lịch sử thông tin hợp đồng cũ."
    )
    add_paragraph_tnr(doc,
        "Đồ án “Quản lý phòng trọ” được xây dựng nhằm tin học hóa và tự động hóa toàn bộ các hoạt động nghiệp vụ nêu trên bằng một ứng dụng desktop chuyên nghiệp "
        "viết trên nền tảng Java Swing. Hệ thống sử dụng hệ quản trị cơ sở dữ liệu SQL Server để lưu trữ tập trung dữ liệu một cách an toàn, JDBC làm cầu nối trung gian "
        "truy cập cơ sở dữ liệu, FlatLaf để tạo giao diện Look and Feel hiện đại và JFreeChart hỗ trợ hiển thị biểu đồ thống kê doanh số trực quan. Báo cáo đồ án môn học này "
        "sẽ trình bày chi tiết về lý do hình thành đề tài, phân tích yêu cầu bài toán, thiết kế kiến trúc phần mềm, cấu trúc cơ sở dữ liệu, các chức năng nghiệp vụ "
        "đã hoàn thành, tự đánh giá dự án và định hướng phát triển trong tương lai."
    )

    add_heading(doc, "MỤC LỤC", 1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_toc(p)
    doc.add_page_break()


def body(doc):
    # ------------------ CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI ------------------
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    
    add_heading(doc, "1.1. Lý do chọn đề tài (Vấn đề cần giải quyết)", 2)
    add_paragraph_tnr(doc, 
        "Hiện nay, nhu cầu thuê và quản lý phòng trọ tại các đô thị lớn tăng rất nhanh. "
        "Tuy nhiên, việc quản lý truyền thống bằng sổ sách hoặc bảng tính Excel bộc lộ nhiều hạn chế lớn như: "
        "dễ nhầm lẫn khi tính chỉ số điện nước hàng tháng; mất nhiều thời gian tra cứu hồ sơ khách thuê và hợp đồng; "
        "khó theo dõi công nợ, hóa đơn quá hạn; và không có một kênh trao đổi thông tin hiệu quả giữa chủ trọ với người thuê. "
        "Hơn nữa, các chủ trọ lớn sở hữu nhiều dãy trọ khác nhau rất khó theo dõi tổng doanh thu, lợi nhuận thực tế "
        "nếu không có sự hỗ trợ của một phần mềm chuyên nghiệp."
    )
    add_paragraph_tnr(doc,
        "Để giải quyết các vấn đề trên, đồ án nghiên cứu xây dựng phần mềm quản lý phòng trọ nhằm cung cấp một giải pháp "
        "tin học hóa toàn diện. Phần mềm sẽ thay thế các phương thức ghi chép thủ công bằng một ứng dụng desktop trực quan, "
        "giúp tự động hóa các khâu liên quan từ quản lý phòng, đăng ký khách thuê, lập hợp đồng, chốt chỉ số điện nước, "
        "tự động tính hóa đơn và kết xuất thống kê báo cáo doanh thu."
    )
    
    add_heading(doc, "1.2. Ý tưởng giải quyết và Phương pháp tiếp cận", 2)
    add_paragraph_tnr(doc,
        "Đề tài hướng tới việc xây dựng một hệ thống phần mềm quản lý đa chủ hộ (multi-tenancy) trên nền tảng Java Swing và cơ sở dữ liệu SQL Server. "
        "Ý tưởng giải quyết cốt lõi bao gồm:"
    )
    add_bullets(doc, [
        "Phân tách dữ liệu theo phiên đăng nhập (Session): Dữ liệu của từng chủ trọ được phân biệt bằng tên đăng nhập (TenDN) hoặc mã chủ trọ (MaChuTro), đảm bảo tính riêng tư, bảo mật dữ liệu tuyệt đối giữa các tài khoản chủ trọ khác nhau.",
        "Liên kết nghiệp vụ tự động: Hệ thống tự động đồng bộ hóa trạng thái phòng trọ khi có khách thuê đến hoặc rời đi, tự động tạo các bản ghi dịch vụ và chỉ số điện nước mặc định cho tháng hiện tại, giảm thiểu tối đa các bước thao tác nhập liệu thủ công.",
        "Thiết kế giao diện dashboard tập trung: Giúp chủ trọ có cái nhìn tổng quan ngay lập tức về số lượng phòng đang thuê, phòng trống, số lượng khách thuê và tổng hợp biểu đồ cột về doanh thu theo tháng."
    ])
    
    add_heading(doc, "1.3. Mục tiêu của hệ thống", 2)
    add_bullets(doc, [
        "Xây dựng phần mềm giao diện desktop thân thiện, hiện đại sử dụng thư viện FlatLaf.",
        "Tự động tính toán tiền phòng, tiền điện nước và tiền dịch vụ dựa trên chỉ số đầu/cuối của tháng và đơn giá cấu hình sẵn.",
        "Hỗ trợ phân quyền người dùng rõ ràng giữa Chủ trọ (quyền quản trị, chỉnh sửa) và Khách thuê (quyền xem thông tin dịch vụ, hóa đơn cá nhân).",
        "Trực quan hóa doanh thu thông qua biểu đồ cột JFreeChart phục vụ báo cáo tài chính định kỳ.",
        "Tự động thực hiện migration (cập nhật cấu trúc database) và đồng bộ dữ liệu khi khởi chạy ứng dụng."
    ])
    
    add_heading(doc, "1.4. Phạm vi thực hiện", 2)
    add_paragraph_tnr(doc,
        "Đồ án tập trung nghiên cứu và triển khai ứng dụng dạng Desktop App dành cho môi trường nội bộ. "
        "Hệ thống đã xây dựng xong cấu trúc cơ sở dữ liệu với 7 bảng nghiệp vụ liên kết chặt chẽ. "
        "Hệ thống chạy ổn định trên môi trường Windows thông qua kết nối SQL Server cục bộ. "
        "Trong phạm vi hiện tại, đồ án chưa tích hợp thanh toán trực tuyến qua ngân hàng/ví điện tử hoặc cơ chế gửi hóa đơn tự động qua email."
    )

    # ------------------ CHƯƠNG 2. CÔNG NGHỆ VÀ MÔI TRƯỜNG PHÁT TRIỂN ------------------
    add_heading(doc, "CHƯƠNG 2. CÔNG NGHỆ VÀ MÔI TRƯỜNG PHÁT TRIỂN", 1)
    add_paragraph_tnr(doc,
        "Hệ thống được phát triển sử dụng các công nghệ, thư viện mã nguồn mở và môi trường cấu hình chuẩn học thuật, "
        "đảm bảo tính ổn định và khả năng tương thích cao. Chi tiết cấu hình được tổng hợp tại bảng dưới đây:"
    )
    add_kv_table(doc, [
        ("Ngôn ngữ lập trình", "Java, project cấu hình trên Eclipse với JRE JavaSE-21."),
        ("Giao diện người dùng", "Java Swing, bố cục linh hoạt theo JFrame/JPanel/JDialog và JTabbedPane."),
        ("Cơ sở dữ liệu", "Microsoft SQL Server, database QuanLyPhongTro."),
        ("Kết nối cơ sở dữ liệu", "Microsoft JDBC Driver mssql-jdbc-13.2.1.jre11.jar."),
        ("Thư viện giao diện", "FlatLaf 3.2.5 hỗ trợ Look & Feel hiện đại, chuyển đổi giao diện sáng/tối."),
        ("Thư viện biểu đồ", "JFreeChart 1.5.3 và JCommon 1.0.23 dùng cho dashboard và thống kê doanh thu."),
        ("IDE phát triển", "Eclipse IDE Java project, mã nguồn đặt trong src và output build tại bin.")
    ])
    add_caption(doc, "Bảng 2.1. Danh mục công nghệ và thư viện sử dụng trong đồ án môn học.")

    # ------------------ CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG ------------------
    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ CHI TIẾT HỆ THỐNG", 1)
    
    add_heading(doc, "3.1. Các tác nhân sử dụng hệ thống (Actors)", 2)
    add_kv_table(doc, [
        ("Chủ trọ (Landlord)", "Tác nhân chính, có quyền tối cao quản trị toàn bộ hệ thống phòng trọ, thông tin khách thuê, hợp đồng, dịch vụ, điện nước, hóa đơn và xem báo cáo tài chính doanh thu."),
        ("Khách thuê (Tenant)", "Đăng nhập bằng tài khoản được liên kết để tự xem thông tin phòng của mình, biểu phí dịch vụ, chỉ số điện nước tiêu thụ và trạng thái hóa đơn hàng tháng."),
        ("Hệ thống (System)", "Tự động chạy script cập nhật cơ sở dữ liệu (Migration), kiểm soát phiên đăng nhập (Session) và tự động đồng bộ hóa hóa đơn khi có thay đổi chỉ số dịch vụ.")
    ])
    add_caption(doc, "Bảng 3.1. Phân tích quyền hạn của các tác nhân trong hệ thống.")

    add_heading(doc, "3.2. Sơ đồ Use Case tổng quát", 2)
    add_paragraph_tnr(doc,
        "Sơ đồ Use Case tổng quát dưới đây mô phỏng các mối tương quan nghiệp vụ chức năng chính của hệ thống và sự tương tác trực quan "
        "giữa hai tác nhân chính (Chủ trọ và Khách thuê) với các chức năng tương ứng:"
    )
    
    # Insert Use Case diagram
    add_picture(doc, "usecase_diagram.png", 5.5)
    add_caption(doc, "Hình 3.1. Sơ đồ Use Case tổng quát của hệ thống phần mềm.")
    
    add_paragraph_tnr(doc,
        "Thuyết minh chi tiết sơ đồ Use Case: "
        "Chủ trọ đóng vai trò quản trị viên tối cao của hệ thống. Để làm việc, chủ trọ cần đăng ký tài khoản quản trị và đăng nhập thành công. "
        "Chủ trọ có toàn quyền thực hiện các Use Case nghiệp vụ quan trọng bao gồm: Quản lý thông tin phòng trọ, Quản lý hồ sơ khách thuê và tạo liên kết tài khoản "
        "cho khách, Lập và theo dõi hợp đồng thuê phòng, Nhập chỉ số và tự động tính tiền điện nước hàng tháng, Cấu hình biểu phí dịch vụ phòng, Lập hóa đơn và kiểm soát "
        "trạng thái thanh toán của từng phòng. Ngoài ra, chủ trọ có thể theo dõi nhanh tổng hợp tình hình kinh doanh thông qua màn hình Dashboard. "
        "Đối với tác nhân Khách thuê, hệ thống phân quyền ở mức thấp hơn, chỉ cho phép thực hiện Use Case đăng nhập bằng tài khoản được cấp để xem thông tin chi tiết "
        "về chỉ số dịch vụ tiêu thụ, biểu phí của phòng mình và trạng thái thanh toán hóa đơn cá nhân hàng tháng."
    )

    add_heading(doc, "3.3. Yêu cầu chức năng chi tiết", 2)
    add_bullets(doc, [
        "Đăng nhập & Đăng ký: Xác thực tài khoản của Chủ trọ và Khách thuê, tự động định tuyến màn hình làm việc tương ứng sau khi đăng nhập thành công.",
        "Quản lý phòng trọ: Quản lý danh sách phòng gồm mã phòng, tầng, diện tích, giá phòng, trạng thái phòng (Trống/Đang thuê).",
        "Quản lý khách thuê: Quản lý thông tin họ tên, số điện thoại, số căn cước công dân (CCCD), quê quán, phòng đang ở và tài khoản liên kết đăng nhập.",
        "Quản lý hợp đồng: Lưu trữ thông tin hợp đồng thuê phòng giữa chủ trọ và khách hàng (ngày bắt đầu, ngày kết thúc, tiền đặt cọc và trạng thái hiệu lực).",
        "Quản lý điện nước: Ghi nhận chỉ số điện (DienDau, DienCuoi), chỉ số nước (NuocDau, NuocCuoi) theo từng tháng, tự động tính tiền điện nước theo đơn giá.",
        "Quản lý dịch vụ: Cấu hình các dịch vụ như Wifi, thang máy, phí giữ xe, vệ sinh/PCCC theo từng phòng.",
        "Quản lý hóa đơn: Tự động tính toán tổng hóa đơn hàng tháng gồm (Tiền phòng + Tiền điện + Tiền nước + Tiền dịch vụ) và cập nhật trạng thái thu tiền (Đã thu/Chưa thu).",
        "Dashboard và Thống kê: Hiển thị nhanh số phòng trống, số khách thuê, số hợp đồng còn hiệu lực và vẽ biểu đồ doanh thu hàng tháng thông qua JFreeChart."
    ])

    add_heading(doc, "3.4. Yêu cầu phi chức năng", 2)
    add_bullets(doc, [
        "Hiệu năng: Các câu lệnh truy vấn dữ liệu SQL Server được tối ưu hóa bằng PreparedStatement để tăng tốc độ phản hồi dưới 1 giây.",
        "Tính an toàn dữ liệu: Áp dụng cơ chế lọc dữ liệu multi-tenancy thông qua câu lệnh SQL WHERE TenDN = ?. Mỗi chủ trọ chỉ có quyền thao tác trên dữ liệu thuộc quyền sở hữu của mình, tránh rò rỉ dữ liệu.",
        "Đồng bộ hóa giao diện: Sử dụng mô hình DataRefreshEvent để tự động cập nhật dữ liệu trên tất cả các tab giao diện ngay khi có thay đổi (Thêm/Sửa/Xóa) ở một màn hình bất kỳ.",
        "Trải nghiệm người dùng: Bố cục trực quan, sử dụng màu sắc đặc trưng để nhận diện nhanh các phòng trống (màu xanh lá) và phòng đang thuê (màu đỏ/vàng)."
    ])

    # ------------------ CHƯƠNG 4. THIẾT KẾ KIẾN TRÚC PHẦN MỀM ------------------
    add_heading(doc, "CHƯƠNG 4. THIẾT KẾ KIẾN TRÚC PHẦN MỀM", 1)
    add_paragraph_tnr(doc,
        "Để đảm bảo tính modular, dễ bảo trì và mở rộng trong tương lai, đồ án tổ chức mã nguồn ứng dụng Java Swing theo mô hình kiến trúc phân lớp "
        "(3-Tier Architecture) giúp phân tách rõ ràng trách nhiệm giữa phần giao diện hiển thị, logic nghiệp vụ dữ liệu và các thực thể lưu trữ."
    )
    
    # Insert Architecture diagram
    add_picture(doc, "architecture_diagram.png", 5.5)
    add_caption(doc, "Hình 4.1. Sơ đồ kiến trúc phân lớp của hệ thống phần mềm.")
    
    add_paragraph_tnr(doc,
        "Thuyết minh chi tiết sơ đồ kiến trúc: "
        "Mã nguồn dự án được tổ chức chặt chẽ thành 4 thành phần chính bao gồm:"
    )
    add_bullets(doc, [
        "Lớp Giao Diện (Presentation Layer - ui): Chứa các lớp hiển thị giao diện người dùng Swing như MainFrame, KhachThueFrame, LoginForm, RegisterForm và các Panel con chuyên trách nghiệp vụ (PhongTroPanel, KhachThuePanel...). Tầng UI chịu trách nhiệm tiếp nhận thao tác click chuột, nhập liệu của người dùng, hiển thị thông tin dạng bảng biểu và gọi xử lý từ tầng DAO.",
        "Lớp Truy Xuất Dữ Liệu (Data Access Layer - dao): Gồm các lớp DAO tương ứng với từng thực thể (TaiKhoanDAO, PhongTroDAO, KhachThueDAO, HopDongDAO, DienNuocDAO, DichVuDAO, HoaDonDAO, ThongKeDAO). Đây là nơi trực tiếp thiết lập kết nối, soạn thảo các câu lệnh SQL với các tham số an toàn (PreparedStatement) để truy vấn tới cơ sở dữ liệu SQL Server.",
        "Lớp Thực Thể (Model Layer - model): Khai báo cấu trúc lớp dữ liệu thuần túy (TaiKhoan.java, PhongTro.java, KhachThue.java...) để đóng gói dữ liệu lấy từ SQL Server và truyền tải dữ liệu an toàn qua lại giữa lớp giao diện UI và lớp DAO.",
        "Lớp Tiện Ích (Util Layer): Cung cấp các công cụ dùng chung như thiết lập driver kết nối (DatabaseConnection), đồng bộ nâng cấp schema tự động (DatabaseMigration) và quản lý đối tượng người dùng đang làm việc trong phiên hiện tại (Session)."
    ])

    # ------------------ CHƯƠNG 5. THIẾT KẾ CƠ SỞ DỮ LIỆU ------------------
    add_heading(doc, "CHƯƠNG 5. THIẾT KẾ CƠ SỞ DỮ LIỆU", 1)
    
    add_heading(doc, "5.1. Sơ đồ liên kết thực thể (ERD)", 2)
    add_paragraph_tnr(doc,
        "Cơ sở dữ liệu của phần mềm Quản lý phòng trọ được thiết kế chuẩn hóa ở dạng chuẩn 3NF, bao gồm 7 bảng dữ liệu liên kết chặt chẽ thông qua các khóa chính "
        "và khóa ngoại để đảm bảo tính toàn vẹn dữ liệu nghiệp vụ và hỗ trợ mô hình đa tài khoản chủ trọ (multi-tenancy):"
    )
    
    # Insert ERD diagram
    add_picture(doc, "erd_diagram.png", 5.8)
    add_caption(doc, "Hình 5.1. Sơ đồ thực thể liên kết (ERD) của cơ sở dữ liệu.")
    
    add_paragraph_tnr(doc,
        "Thuyết minh chi tiết sơ đồ ERD và các mối quan hệ khóa ngoại:"
    )
    add_bullets(doc, [
        "Bảng TaiKhoan đóng vai trò lưu trữ tài khoản đăng nhập. Mỗi chủ trọ có một tài khoản xác định qua khóa chính TenDN. Mỗi tài khoản chủ trọ sở hữu nhiều phòng trọ trong bảng PhongTro liên kết qua trường TenDN (mối quan hệ 1 - Nhiều).",
        "Bảng PhongTro lưu thông tin chi tiết các phòng. Khóa chính của bảng là khóa kết hợp gồm (MaPhong, TenDN) để giải quyết bài toán trùng mã phòng giữa các chủ trọ khác nhau. Một phòng trọ (PhongTro) liên kết với nhiều khách thuê trong bảng KhachThue thông qua khóa ngoại MaPhong.",
        "Bảng KhachThue lưu trữ danh sách khách thuê và liên kết tài khoản riêng của khách. Trường MaChuTro là khóa ngoại trỏ tới bảng TaiKhoan để phân chia quyền sở hữu khách hàng cho từng chủ trọ cụ thể.",
        "Bảng HopDong quản lý hợp đồng thuê phòng giữa chủ trọ và khách thuê. Mỗi phòng trọ phát sinh nhiều hợp đồng theo thời gian (mối quan hệ 1 - Nhiều) thông qua trường MaPhong.",
        "Bảng DienNuoc ghi nhận chỉ số điện nước tiêu thụ hàng tháng của từng phòng trọ. Trường MaPhong là khóa ngoại liên kết tới bảng PhongTro.",
        "Bảng DichVu lưu cấu hình phí dịch vụ cố định (Wifi, xe, vệ sinh...) riêng biệt cho từng phòng trọ thông qua khóa ngoại MaPhong.",
        "Bảng HoaDon lưu trữ lịch sử hóa đơn thanh toán tiền phòng phát sinh hàng tháng. Trường TenPhong (tương ứng với MaPhong) đóng vai trò khóa ngoại liên kết với bảng PhongTro để truy xuất đơn giá phòng."
    ])

    add_heading(doc, "5.2. Mô tả cấu trúc các bảng chi tiết", 2)
    add_paragraph_tnr(doc,
        "Dưới đây là mô tả chi tiết kiểu dữ liệu, ràng buộc và vai trò nghiệp vụ của 7 bảng dữ liệu chính được sử dụng trực tiếp trong hệ thống phần mềm:"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Bảng dữ liệu", "Vai trò nghiệp vụ", "Các thuộc tính tiêu biểu"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    rows = [
        ("TaiKhoan", "Lưu trữ thông tin tài khoản đăng nhập hệ thống và phân quyền vai trò.", "TenDN, MatKhau, Quyen, Email"),
        ("PhongTro", "Lưu trữ danh sách các phòng trọ của từng chủ trọ sở hữu.", "MaPhong, Tang, TrangThai, GiaPhong, DienTich, TenDN"),
        ("KhachThue", "Lưu trữ lý lịch khách hàng thuê phòng và liên kết tài khoản đăng nhập.", "MaKhach, HoTen, CCCD, SDT, QueQuan, MaPhong, TrangThai, TenDN, MaChuTro"),
        ("HopDong", "Lưu thông tin hợp đồng pháp lý, tiền đặt cọc và thời hạn thuê phòng.", "MaHD, TenKhachHang, MaPhong, NgayBatDau, NgayKetThuc, TienCoc, TrangThai, TenDN"),
        ("DienNuoc", "Lưu trữ chỉ số điện nước đầu và cuối theo từng tháng tiêu dùng.", "ID, MaPhong, ThangNam, DienDau, DienCuoi, NuocDau, NuocCuoi, TienDien, TienNuoc, TenDN"),
        ("DichVu", "Lưu cấu hình chi phí các dịch vụ cố định đi kèm theo từng phòng.", "MaDV, MaPhong, Wifi, ThangMay, GiuXe, PCCC_VeSinh, TongTienDV, TenDN"),
        ("HoaDon", "Lưu trữ hóa đơn thanh toán tiền phòng, điện nước và dịch vụ hàng tháng.", "MaHoaDon, TenPhong, NgayLap, TienPhong, TienDien, TienNuoc, TienDichVu, TrangThai, TenDN"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0))
    style_table(table, [1.3, 2.4, 2.7])
    add_caption(doc, "Bảng 5.3. Chi tiết cấu trúc dữ liệu và vai trò các bảng trong database.")

    add_heading(doc, "5.3. Cơ chế Migration và đồng bộ dữ liệu tự động", 2)
    add_paragraph_tnr(doc,
        "Khi ứng dụng khởi chạy, lớp DatabaseMigration sẽ tự động kiểm tra sự tồn tại của các cột nghiệp vụ trong database. "
        "Nếu phát hiện các cột mới phục vụ multi-tenancy chưa được tạo (ví dụ TenDN, MaChuTro), hệ thống sẽ tự động gửi lệnh ALTER TABLE để thêm cột và gán giá trị mặc định dựa trên dữ liệu hiện có. "
        "Ngoài ra, hệ thống tự động kiểm tra và khởi tạo bản ghi dịch vụ (DichVu), chỉ số điện nước (DienNuoc) và hóa đơn (HoaDon) cho tháng hiện tại đối với tất cả những phòng trọ đang có khách thuê ở trạng thái 'Đang ở'. "
        "Điều này giúp dữ liệu cũ tương thích hoàn toàn với cấu trúc phần mềm mới một cách trơn tru mà không cần can thiệp thủ công từ hệ trị quản cơ sở dữ liệu."
    )

    # ------------------ CHƯƠNG 6. MÔ TẢ CÁC CHỨC NĂNG CHÍNH ------------------
    add_heading(doc, "CHƯƠNG 6. MÔ TẢ CÁC CHỨC NĂNG CHÍNH", 1)
    add_paragraph_tnr(doc,
        "Các chức năng chính của hệ thống được lập trình phân tách thành các Module giao diện (Panel) tương ứng với các lớp xử lý DAO chuyên biệt:"
    )
    
    function_rows = [
        ("Đăng nhập/đăng ký", "LoginForm, RegisterForm, TaiKhoanDAO", "Xác thực tài khoản, mã hóa phân quyền giao diện hiển thị cho chủ trọ hoặc khách thuê."),
        ("Dashboard", "DashboardPanel, HoaDonDAO, PhongTroDAO", "Hiển thị trực quan tổng số phòng, phòng trống/đang thuê và biểu đồ cột doanh thu hàng tháng."),
        ("Thống kê doanh thu", "ThongKePanel, ThongKeDAO", "Lọc doanh thu thực tế đã thu, tổng tiền nợ chưa thu, số lượng khách thuê và số hợp đồng hiệu lực."),
        ("Quản lý phòng trọ", "PhongTroPanel, PhongTroDialog, PhongTroDAO", "Thêm, sửa, xóa phòng trọ, hiển thị phòng dạng sơ đồ trực quan với các màu sắc biểu thị trạng thái."),
        ("Quản lý khách thuê", "KhachThuePanel, KhachThueDialog, KhachThueDAO", "Quản lý thông tin khách thuê, kiểm tra trùng căn cước công dân (CCCD), cập nhật trạng thái phòng khi khách dọn vào/rời đi."),
        ("Quản lý hợp đồng", "HopDongPanel, HopDongDAO", "Tạo lập hợp đồng thuê phòng với các điều khoản đặt cọc và ngày bắt đầu/kết thúc."),
        ("Quản lý điện nước", "DienNuocPanel, DienNuocDAO", "Nhập chỉ số điện nước cuối tháng, tự động tính tiền chênh lệch và cập nhật sang hóa đơn chưa thu."),
        ("Quản lý dịch vụ", "DichVuPanel, DichVuDAO", "Cấu hình chi tiết các khoản phí dịch vụ (Wifi, thang máy, gửi xe, vệ sinh) áp dụng theo từng phòng."),
        ("Quản lý hóa đơn", "HoaDonPanel, HoaDonDAO", "Tự động kết xuất hóa đơn hàng tháng dựa trên tiền phòng, tiền điện nước và tiền dịch vụ. Cho phép đánh dấu trạng thái thu tiền."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Phân hệ chức năng", "Lớp giao diện & xử lý", "Nhiệm vụ chi tiết"]):
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in function_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=(i == 0))
    style_table(table, [1.5, 2.2, 2.7])
    add_caption(doc, "Bảng 6.1. Chi tiết phân hệ chức năng chính của hệ thống.")

    add_heading(doc, "6.1. Liên kết nghiệp vụ tự động hóa nổi bật", 2)
    add_paragraph_tnr(doc,
        "Một trong những ưu điểm lớn nhất của phần mềm là cơ chế tự động liên kết dữ liệu giữa các nghiệp vụ khác nhau:"
    )
    add_bullets(doc, [
        "Đồng bộ trạng thái phòng: Khi thêm mới một khách thuê ở trạng thái 'Đang ở', hệ thống tự động đổi trạng thái phòng tương ứng từ 'Trống' sang 'Đang thuê'. Ngược lại, khi khách rời phòng hoặc bị xóa khỏi hệ thống, trạng thái phòng lập tức chuyển về 'Trống'.",
        "Tạo hợp đồng mặc định: Khi khách hàng thuê phòng, hệ thống hỗ trợ tự động tạo một hợp đồng thuê mặc định có thời hạn 1 năm với giá trị tiền cọc bằng đúng giá phòng trọ hiện tại.",
        "Đồng bộ hóa chỉ số tiêu thụ: Khi chủ trọ nhập chỉ số điện nước cuối tháng trên DienNuocPanel, tiền điện (đơn giá 3.500đ/kWh) và tiền nước (đơn giá 15.000đ/m3) được tự động tính toán và cập nhật thẳng vào hóa đơn chưa thanh toán của phòng đó trong tháng.",
        "Khởi tạo dữ liệu tháng mới: Hệ thống tự động phát hiện và sinh bản ghi điện nước đầu kỳ, dịch vụ cố định và hóa đơn chờ thu khi bước sang chu kỳ tháng tiếp theo cho các phòng đang thuê."
    ])

    # ------------------ CHƯƠNG 7. ĐÁNH GIÁ KẾT QUẢ ------------------
    add_heading(doc, "CHƯƠNG 7. TỰ ĐÁNH GIÁ VÀ ĐÁNH GIÁ KẾT QUẢ", 1)
    
    add_heading(doc, "7.1. Trả lời các câu hỏi cốt lõi của đề tài", 2)
    
    add_heading(doc, "a) Vấn đề cần giải quyết trong đề tài của mình là gì?", 3)
    add_paragraph_tnr(doc,
        "Vấn đề cốt lõi cần giải quyết là sự thiếu hiệu quả, dễ nhầm lẫn và tốn thời gian của quy trình quản lý phòng trọ thủ công bằng sổ sách và Excel. "
        "Đặc biệt là việc theo dõi riêng biệt chỉ số điện nước hàng tháng, tính toán hóa đơn phức tạp bao gồm nhiều khoản phí dịch vụ khác nhau, "
        "quản lý thời hạn và tiền cọc hợp đồng, kiểm soát công nợ chưa đóng, và bảo mật thông tin giữa nhiều chủ trọ khác nhau sở hữu các khu trọ riêng biệt."
    )
    
    add_heading(doc, "b) Cách giải quyết các vấn đề đó như thế nào?", 3)
    add_paragraph_tnr(doc,
        "Đề tài đã giải quyết các vấn đề trên bằng cách xây dựng một ứng dụng desktop Java Swing hoàn chỉnh kết nối cơ sở dữ liệu SQL Server. "
        "Các giải pháp kỹ thuật cụ thể gồm: thiết kế cấu trúc CSDL chuẩn hóa gồm 7 bảng nghiệp vụ liên kết chặt chẽ; tích hợp cơ chế lọc thông tin theo tài khoản "
        "đăng nhập (WHERE TenDN = ?) để hỗ trợ đa chủ trọ bảo mật dữ liệu; tự động hóa tính toán tiền điện nước và tiền dịch vụ dựa trên chỉ số đầu/cuối của tháng; "
        "tự động cập nhật trạng thái phòng trọ trống/thuê khi quản lý khách thuê; và cung cấp biểu đồ JFreeChart để trực quan hóa doanh thu tài chính."
    )
    
    add_heading(doc, "c) Kết quả đã làm được là gì?", 3)
    add_paragraph_tnr(doc,
        "Đồ án đã xây dựng thành công phần mềm Quản lý phòng trọ hoạt động ổn định trên nền tảng Windows với đầy đủ các phân hệ chức năng: "
        "Đăng nhập & Đăng ký tài khoản; Dashboard trực quan hóa trạng thái phòng và biểu đồ doanh thu; Phân hệ quản lý phòng trọ; Phân hệ quản lý khách thuê; "
        "Phân hệ quản lý hợp đồng thuê phòng; Phân hệ quản lý chỉ số điện nước tiêu thụ; Phân hệ cấu hình phí dịch vụ; Phân hệ tự động tính toán hóa đơn tháng. "
        "Hệ thống cũng đã xây dựng cơ chế tự động Migration cập nhật schema dữ liệu khi nâng cấp ứng dụng."
    )
    
    add_heading(doc, "d) Có thỏa mãn yêu cầu đề ra không?", 3)
    add_paragraph_tnr(doc,
        "Kết quả đạt được hoàn toàn thỏa mãn các yêu cầu đề ra ban đầu của đồ án môn học. Phần mềm giải quyết triệt để tất cả các khó khăn trong nghiệp vụ quản lý phòng trọ thực tế, "
        "đáp ứng các tiêu chí về tính an toàn dữ liệu, tính chính xác trong tính toán tài chính điện nước, tính trực quan trong báo cáo biểu đồ, "
        "và mang lại trải nghiệm người dùng hiện đại và mượt mà nhờ thư viện FlatLaf."
    )

    add_heading(doc, "7.2. Các ưu điểm nổi bật của phần mềm", 2)
    add_bullets(doc, [
        "Xây dựng thành công giải pháp multi-tenancy phân tách dữ liệu an toàn cho nhiều chủ trọ trên cùng một cơ sở dữ liệu.",
        "Mã nguồn được tổ chức sạch sẽ theo cấu trúc phân lớp rõ ràng (UI - DAO - Model - Util) giúp dễ bảo trì và mở rộng.",
        "Cơ chế tự động hóa nghiệp vụ liên kết chặt chẽ giữa khách thuê, trạng thái phòng trọ, điện nước tiêu thụ và hóa đơn hàng tháng.",
        "Sử dụng PreparedStatement để ngăn chặn các lỗi chèn mã độc SQL Injection và tăng hiệu năng truy vấn dữ liệu.",
        "Giao diện hiện đại, trực quan, hỗ trợ hiển thị biểu đồ thống kê doanh thu tài chính chuyên nghiệp."
    ])

    add_heading(doc, "7.3. Những điểm hạn chế còn tồn tại", 2)
    add_bullets(doc, [
        "Cấu hình thông tin kết nối cơ sở dữ liệu vẫn đang được khai báo trực tiếp trong mã nguồn Java, chưa được tách ra tệp cấu hình bên ngoài.",
        "Mật khẩu tài khoản của người dùng lưu trữ dưới dạng văn bản thuần túy (Plaintext), chưa được mã hóa băm bảo mật (như MD5/BCrypt).",
        "Chưa triển khai chức năng xuất hóa đơn hàng tháng ra tệp PDF hoặc gửi trực tiếp thông báo hóa đơn qua Email/Zalo của khách thuê."
    ])

    add_heading(doc, "7.4. Hướng phát triển trong tương lai", 2)
    add_bullets(doc, [
        "Tách cấu hình kết nối database ra tệp properties bên ngoài hoặc biến môi trường hệ thống để dễ dàng cài đặt khi triển khai thực tế.",
        "Áp dụng giải thuật băm mật khẩu BCrypt để nâng cao tính bảo mật cho tài khoản người dùng đăng nhập.",
        "Bổ sung phân hệ xuất hóa đơn sang định dạng tệp PDF/Excel phục vụ in ấn hóa đơn và tích hợp cổng gửi email thông báo hóa đơn tự động."
    ])

    # ------------------ KẾT LUẬN ------------------
    add_heading(doc, "KẾT LUẬN", 1)
    add_paragraph_tnr(doc,
        "Đồ án môn học “Phần mềm quản lý phòng trọ” đã xây dựng hoàn thiện một ứng dụng desktop đáp ứng toàn diện các nghiệp vụ quản lý phòng trọ trong thực tế. "
        "Hệ thống giúp các chủ trọ giảm thiểu tối đa sai sót, tự động hóa quy trình tính toán hóa đơn phức tạp và theo dõi doanh thu một cách trực quan, chính xác. "
        "Mặc dù còn một số hạn chế nhỏ về bảo mật băm mật khẩu hay tách cấu hình, ứng dụng đã cung cấp một nền tảng vững chắc và hoàn toàn khả thi để phát triển nâng cấp "
        "thành một sản phẩm phần mềm thương mại hóa."
    )

    # ------------------ TÀI LIỆU THAM KHẢO ------------------
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_numbered(doc, [
        "Mã nguồn project QuanLyPhongTro trong workspace C:\\Users\\Admin\\eclipse-workspace\\QuanLyPhongTro.",
        "Tài liệu hướng dẫn phát triển giao diện đồ họa Java Swing (JFrame, JPanel, JTable, JTabbedPane, JDialog).",
        "Tài liệu đặc tả kết nối cơ sở dữ liệu Microsoft JDBC Driver for SQL Server.",
        "Tài liệu ngôn ngữ truy vấn cấu trúc SQL Server (SELECT, INSERT, UPDATE, DELETE, JOIN và các ràng buộc dữ liệu).",
        "Tài liệu hướng dẫn sử dụng thư viện JFreeChart, FlatLaf và thư viện docx (python-docx) trong lập trình Python và Java."
    ])

    # ------------------ PHỤ LỤC. DANH SÁCH FILE CHÍNH ------------------
    add_heading(doc, "PHỤ LỤC. DANH SÁCH FILE NGHIỆP VỤ CHÍNH", 1)
    add_bullets(doc, [
        "Main.java: Điểm khởi đầu ứng dụng, gọi DatabaseMigration và thiết lập giao diện đăng nhập.",
        "DatabaseConnection.java: Thiết lập kết nối an toàn với cơ sở dữ liệu SQL Server.",
        "DatabaseMigration.java: Quản lý tự động nâng cấp cấu trúc bảng dữ liệu nghiệp vụ và đồng bộ dữ liệu.",
        "Session.java: Lưu trữ thông tin tài khoản đang làm việc trong phiên hiện tại.",
        "MainFrame.java: Khung làm việc chính chứa các tab chức năng dành cho chủ trọ.",
        "LoginForm.java, RegisterForm.java: Màn hình đăng nhập và đăng ký tài khoản.",
        "Các lớp đối tượng DAO: Thực hiện các câu lệnh SQL truy xuất dữ liệu từ database.",
        "Các lớp đối tượng Model: Biểu diễn dữ liệu thực thể nghiệp vụ dưới dạng đối tượng Java.",
        "Các Panel UI: Thiết kế giao diện hiển thị trực quan cho các phân hệ chức năng."
    ])


def main():
    # 1. Generate diagrams dynamically
    print("Generating system diagrams using Pillow...")
    draw_usecase_diagram("usecase_diagram.png")
    draw_architecture_diagram("architecture_diagram.png")
    draw_erd_diagram("erd_diagram.png")
    
    # 2. Construct document
    print("Constructing report Word document...")
    doc = Document()
    configure_document(doc)
    enable_update_fields(doc)
    cover_page(doc)
    front_matter(doc)
    body(doc)
    
    # 3. Post-process to clear indentation from centered elements
    for p in doc.paragraphs:
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            p.paragraph_format.first_line_indent = Inches(0)
            
    doc.core_properties.title = "Báo cáo đồ án Phần mềm quản lý phòng trọ"
    doc.core_properties.subject = "QuanLyPhongTro"
    doc.core_properties.author = "VKU IT Students"
    doc.save(OUT)
    print(f"Report document created successfully at {OUT}!")


if __name__ == "__main__":
    main()
